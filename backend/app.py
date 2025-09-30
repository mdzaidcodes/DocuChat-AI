"""
DocuChat AI - Flask Backend
This backend handles document uploads, vector embeddings, and RAG-based chat interactions.
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import logging
from typing import List, Dict
import chromadb
from chromadb.config import Settings
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.llms import Ollama
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
import uuid
import json
from datetime import datetime

# Initialize Flask app
app = Flask(__name__)

# Configure CORS to allow requests from Next.js frontend
CORS(app, resources={
    r"/*": {
        "origins": ["http://localhost:3000"],
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type"]
    }
})

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
UPLOAD_FOLDER = 'uploads'
CHROMA_DB_DIR = 'chroma_db'
CHAT_HISTORY_DIR = 'chat_history'
CHAT_HISTORY_FILE = os.path.join(CHAT_HISTORY_DIR, 'history.json')
SESSION_FILE = os.path.join(CHAT_HISTORY_DIR, 'session.json')
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CHROMA_DB_DIR, exist_ok=True)
os.makedirs(CHAT_HISTORY_DIR, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# Global variables for vector store and LLM
vector_store = None
qa_chain = None
document_metadata = {}
all_documents = []  # Store all document chunks for multi-document support
active_session_id = None  # Track which session is currently active

def allowed_file(filename: str) -> bool:
    """Check if the uploaded file has an allowed extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def load_chat_sessions() -> Dict:
    """Load all chat sessions from JSON file."""
    try:
        if os.path.exists(CHAT_HISTORY_FILE):
            with open(CHAT_HISTORY_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # Migrate old format to new format if needed
                if isinstance(data, list):
                    # Old format: convert to sessions
                    logger.info("Migrating old chat history format to sessions")
                    return {"sessions": {}}
                return data
        return {"sessions": {}, "active_session": None}
    except Exception as e:
        logger.error(f"Error loading chat sessions: {str(e)}")
        return {"sessions": {}, "active_session": None}

def save_chat_sessions(data: Dict):
    """Save all chat sessions to JSON file."""
    try:
        with open(CHAT_HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        logger.info("Chat sessions saved successfully")
    except Exception as e:
        logger.error(f"Error saving chat sessions: {str(e)}")

def generate_chat_name(question: str) -> str:
    """Generate a chat name from the first question (max 50 chars)."""
    # Take first question, limit length
    name = question.strip()
    if len(name) > 50:
        name = name[:47] + "..."
    return name

def create_new_session(first_question: str = None) -> str:
    """Create a new chat session and return its ID."""
    global document_metadata
    
    try:
        session_id = str(uuid.uuid4())
        sessions_data = load_chat_sessions()
        
        # Generate name from first question or use default
        if first_question:
            name = generate_chat_name(first_question)
        else:
            name = f"New Chat {len(sessions_data['sessions']) + 1}"
        
        # Store current document metadata with the session (safe copy)
        try:
            docs_snapshot = dict(document_metadata) if document_metadata else {}
        except Exception as e:
            logger.warning(f"Failed to snapshot documents: {str(e)}")
            docs_snapshot = {}
        
        sessions_data['sessions'][session_id] = {
            "id": session_id,
            "name": name,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
            "messages": [],
            "documents": docs_snapshot
        }
        sessions_data['active_session'] = session_id
        save_chat_sessions(sessions_data)
        logger.info(f"✓ Created new session: {session_id} - {name} with {len(docs_snapshot)} documents")
        return session_id
    except Exception as e:
        logger.error(f"Error creating session: {str(e)}")
        raise

def add_message_to_session(session_id: str, question: str, answer: str, citations: List[Dict]):
    """Add a message to a specific chat session."""
    global document_metadata
    
    try:
        sessions_data = load_chat_sessions()
        
        if session_id not in sessions_data['sessions']:
            logger.error(f"❌ Session {session_id} not found")
            return False
        
        message = {
            "id": str(uuid.uuid4()),
            "timestamp": datetime.now().isoformat(),
            "question": question,
            "answer": answer,
            "citations": citations
        }
        
        sessions_data['sessions'][session_id]['messages'].append(message)
        sessions_data['sessions'][session_id]['updated_at'] = datetime.now().isoformat()
        
        # Update session name if this is the first message
        if len(sessions_data['sessions'][session_id]['messages']) == 1:
            new_name = generate_chat_name(question)
            sessions_data['sessions'][session_id]['name'] = new_name
            logger.info(f"✓ Renamed session to: {new_name}")
        
        # Update documents snapshot if not present
        if 'documents' not in sessions_data['sessions'][session_id]:
            try:
                sessions_data['sessions'][session_id]['documents'] = dict(document_metadata) if document_metadata else {}
            except Exception as e:
                logger.warning(f"Failed to add documents to session: {str(e)}")
                sessions_data['sessions'][session_id]['documents'] = {}
        
        save_chat_sessions(sessions_data)
        message_count = len(sessions_data['sessions'][session_id]['messages'])
        logger.info(f"✓ Added message to session {session_id} (total: {message_count} messages)")
        return True
    except Exception as e:
        logger.error(f"❌ Error adding message to session: {str(e)}")
        return False

def save_session_state():
    """Save current session state (document metadata and file paths) to JSON."""
    try:
        session_data = {
            "document_metadata": document_metadata,
            "last_updated": datetime.now().isoformat()
        }
        with open(SESSION_FILE, 'w', encoding='utf-8') as f:
            json.dump(session_data, f, indent=2)
        logger.info("Session state saved")
    except Exception as e:
        logger.error(f"Error saving session state: {str(e)}")

def load_session_state() -> Dict:
    """Load session state from JSON."""
    try:
        if os.path.exists(SESSION_FILE):
            with open(SESSION_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    except Exception as e:
        logger.error(f"Error loading session state: {str(e)}")
        return {}

def restore_session():
    """
    Restore previous session by reloading documents and rebuilding vector store.
    Called on backend startup if session file exists.
    """
    global document_metadata, vector_store, qa_chain, all_documents
    
    try:
        session_data = load_session_state()
        if not session_data or 'document_metadata' not in session_data:
            logger.info("No previous session found")
            return False
        
        document_metadata = session_data['document_metadata']
        
        if not document_metadata:
            logger.info("No documents in previous session")
            return False
        
        logger.info(f"Restoring session with {len(document_metadata)} documents...")
        
        # Reload all documents from uploads folder
        all_chunks = []
        for file_id, metadata in document_metadata.items():
            # Find the file in uploads
            for filename in os.listdir(UPLOAD_FOLDER):
                if filename.startswith(file_id):
                    file_path = os.path.join(UPLOAD_FOLDER, filename)
                    
                    # Load and chunk the document
                    documents = load_document(file_path)
                    for doc in documents:
                        doc.metadata['source'] = metadata['filename']
                        doc.metadata['file_id'] = file_id
                    
                    chunks = chunk_documents(documents)
                    all_chunks.extend(chunks)
                    logger.info(f"Reloaded: {metadata['filename']}")
                    break
        
        if all_chunks:
            # Rebuild vector store
            create_vector_store(all_chunks, collection_name="docuchat_multi", add_to_existing=False)
            initialize_qa_chain()
            logger.info(f"Session restored successfully with {len(all_chunks)} chunks")
            return True
        
        return False
        
    except Exception as e:
        logger.error(f"Error restoring session: {str(e)}")
        return False

def load_document(file_path: str) -> List:
    """
    Load document based on file type (PDF or DOCX).
    Returns a list of document chunks with metadata.
    """
    file_extension = file_path.rsplit('.', 1)[1].lower()
    
    try:
        if file_extension == 'pdf':
            loader = PyPDFLoader(file_path)
            documents = loader.load()
        elif file_extension in ['docx', 'doc']:
            # Try UnstructuredWordDocumentLoader first
            try:
                loader = UnstructuredWordDocumentLoader(file_path)
                documents = loader.load()
            except Exception as e:
                logger.warning(f"UnstructuredWordDocumentLoader failed: {str(e)}, trying python-docx directly")
                # Fallback: Use python-docx directly
                import docx
                doc = docx.Document(file_path)
                text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
                
                from langchain.schema import Document
                documents = [Document(page_content=text, metadata={"source": file_path})]
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        logger.info(f"✓ Loaded {len(documents)} pages/sections from {file_path}")
        return documents
    except Exception as e:
        logger.error(f"❌ Error loading document: {str(e)}")
        raise

def chunk_documents(documents: List, chunk_size: int = 1000, chunk_overlap: int = 200) -> List:
    """
    Split documents into semantic chunks using RecursiveCharacterTextSplitter.
    This ensures better context preservation and retrieval accuracy.
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    chunks = text_splitter.split_documents(documents)
    logger.info(f"Created {len(chunks)} chunks from documents")
    return chunks

def create_vector_store(chunks: List, collection_name: str = "docuchat", add_to_existing: bool = False, session_id: str = None):
    """
    Create or update a session-specific Chroma vector store with Ollama embeddings.
    Each session has its own isolated vector store to keep documents separate.
    Uses local Ollama instance for embedding generation.
    """
    global vector_store, all_documents, active_session_id
    
    try:
        # Initialize Ollama embeddings (using nomic-embed-text for embeddings)
        embeddings = OllamaEmbeddings(
            model="nomic-embed-text",
            base_url="http://localhost:11434"
        )
        
        # Use session-specific collection name if session_id provided
        if session_id:
            collection_name = f"session_{session_id}"
            
        if add_to_existing and vector_store is not None and active_session_id == session_id:
            # Add new documents to existing vector store (same session)
            all_documents.extend(chunks)
            vector_store.add_documents(chunks)
            logger.info(f"Added {len(chunks)} chunks to session {session_id}")
        else:
            # Create new Chroma vector store for this session
            all_documents = chunks
            active_session_id = session_id
            vector_store = Chroma.from_documents(
                documents=chunks,
                embedding=embeddings,
                collection_name=collection_name,
                persist_directory=CHROMA_DB_DIR
            )
            logger.info(f"✓ Vector store created for session {session_id} with {len(chunks)} chunks")
        
        return vector_store
    except Exception as e:
        logger.error(f"Error creating/updating vector store: {str(e)}")
        raise

def load_session_vector_store(session_id: str):
    """
    Load the vector store for a specific session.
    This allows switching between sessions with different documents.
    """
    global vector_store, active_session_id, qa_chain
    
    try:
        # If already loaded for this session, skip
        if active_session_id == session_id and vector_store is not None:
            logger.info(f"Vector store already loaded for session {session_id}")
            return True
        
        # Load vector store for this session
        embeddings = OllamaEmbeddings(
            model="nomic-embed-text",
            base_url="http://localhost:11434"
        )
        
        collection_name = f"session_{session_id}"
        logger.info(f"Loading vector store for session {session_id}")
        
        vector_store = Chroma(
            collection_name=collection_name,
            embedding_function=embeddings,
            persist_directory=CHROMA_DB_DIR
        )
        
        active_session_id = session_id
        
        # Reinitialize QA chain with the loaded vector store
        initialize_qa_chain()
        
        logger.info(f"✓ Vector store loaded for session {session_id}")
        return True
    except Exception as e:
        logger.error(f"Failed to load vector store for session {session_id}: {str(e)}")
        return False

def initialize_qa_chain():
    """
    Initialize the RetrievalQA chain with Ollama LLM.
    This chain handles question answering with source citations.
    """
    global qa_chain, vector_store
    
    if vector_store is None:
        raise ValueError("Vector store not initialized")
    
    try:
        # Initialize Ollama LLM (llama3:8b)
        llm = Ollama(
            model="llama3:8b",
            base_url="http://localhost:11434",
            temperature=0.3  # Lower temperature for more focused responses
        )
        
        # Custom prompt template for better responses with citations
        prompt_template = """You are a helpful AI assistant that answers questions based on the provided context from uploaded documents. 
        Use the following pieces of context to answer the question at the end. 
        If you don't know the answer or if it's not in the context, just say that you don't have enough information to answer, don't try to make up an answer.
        Always be specific and cite the relevant parts of the context in your answer.

Context: {context}

Question: {question}

Answer: Let me help you with that based on the document."""
        
        PROMPT = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )
        
        # Create RetrievalQA chain with source documents return
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vector_store.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 4}  # Retrieve top 4 most relevant chunks
            ),
            return_source_documents=True,
            chain_type_kwargs={"prompt": PROMPT}
        )
        
        logger.info("QA chain initialized successfully")
        return qa_chain
    except Exception as e:
        logger.error(f"Error initializing QA chain: {str(e)}")
        raise

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint to verify backend is running."""
    return jsonify({
        "status": "healthy",
        "message": "DocuChat AI Backend is running"
    }), 200

@app.route('/upload', methods=['POST'])
def upload_document():
    """
    Endpoint to upload a document (PDF or DOCX).
    Processes the document, creates embeddings, and stores in vector database.
    Supports adding to existing collection or creating new one.
    """
    global document_metadata
    
    # Check if file is in request
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    
    # Check if file is selected
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    
    # Validate file type
    if not allowed_file(file.filename):
        return jsonify({"error": "Invalid file type. Only PDF and DOCX files are allowed"}), 400
    
    # Check if this is adding to existing documents
    add_to_existing = request.form.get('add_to_existing', 'false').lower() == 'true'
    
    try:
        # Secure the filename and save
        filename = secure_filename(file.filename)
        file_id = str(uuid.uuid4())
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}_{filename}")
        file.save(file_path)
        
        logger.info(f"File uploaded: {filename} (add_to_existing: {add_to_existing})")
        
        # Load document
        documents = load_document(file_path)
        
        # Add metadata to documents
        for doc in documents:
            doc.metadata['source'] = filename
            doc.metadata['file_id'] = file_id
        
        # Chunk documents with semantic splitting
        chunks = chunk_documents(documents)
        
        # Create or update vector store (no session tracking - simple mode)
        create_vector_store(chunks, collection_name="docuchat", add_to_existing=add_to_existing, session_id=None)
        
        # Initialize or update QA chain
        initialize_qa_chain()
        
        # Store document metadata
        document_metadata[file_id] = {
            "filename": filename,
            "chunks": len(chunks),
            "pages": len(documents)
        }
        
        # Save session state for persistence
        save_session_state()
        
        total_documents = len(document_metadata)
        
        return jsonify({
            "message": "Document uploaded and processed successfully" if not add_to_existing else "Document added to collection",
            "file_id": file_id,
            "filename": filename,
            "chunks": len(chunks),
            "pages": len(documents),
            "total_documents": total_documents
        }), 200
        
    except Exception as e:
        logger.error(f"Error processing upload: {str(e)}")
        return jsonify({"error": f"Failed to process document: {str(e)}"}), 500

def format_citation_text(text: str, max_sentences: int = 3) -> str:
    """
    Format citation text to show only 2-3 sentences.
    Truncates long text and adds ellipsis if needed.
    """
    import re
    
    # Split into sentences (basic approach)
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    
    # Take first max_sentences
    selected = sentences[:max_sentences]
    result = ' '.join(selected)
    
    # Add ellipsis if we truncated
    if len(sentences) > max_sentences:
        result += '...'
    
    # Limit total length as backup
    if len(result) > 300:
        result = result[:297] + '...'
    
    return result

@app.route('/chat', methods=['POST'])
def chat():
    """
    Endpoint to send a question and receive an AI-generated answer with citations.
    Returns the answer along with source documents for citation.
    Automatically saves conversation to active chat session.
    Citations are formatted as 2-3 sentences with page numbers.
    """
    global qa_chain
    
    # Check if QA chain is initialized
    if qa_chain is None:
        return jsonify({"error": "No document uploaded yet. Please upload a document first."}), 400
    
    # Get question and session_id from request
    data = request.get_json()
    if not data or 'question' not in data:
        return jsonify({"error": "No question provided"}), 400
    
    question = data['question']
    
    try:
        # Get answer from QA chain (using global vector store, no sessions)
        result = qa_chain({"query": question})
        
        # Extract answer and source documents
        answer = result['result']
        source_docs = result['source_documents']
        
        # Format source citations (shortened to 2-3 sentences)
        citations = []
        for i, doc in enumerate(source_docs):
            formatted_text = format_citation_text(doc.page_content, max_sentences=3)
            citations.append({
                "id": i + 1,
                "content": formatted_text,
                "source": doc.metadata.get('source', 'Unknown'),
                "page": doc.metadata.get('page', 'N/A')
            })
        
        logger.info(f"Question answered: {question[:50]}...")
        
        return jsonify({
            "answer": answer,
            "citations": citations,
            "question": question
        }), 200
        
    except Exception as e:
        logger.error(f"Error processing chat: {str(e)}")
        return jsonify({"error": f"Failed to process question: {str(e)}"}), 500

@app.route('/documents', methods=['GET'])
def get_documents():
    """Endpoint to retrieve list of uploaded documents."""
    return jsonify({
        "documents": document_metadata
    }), 200

@app.route('/documents/<file_id>', methods=['DELETE'])
def remove_document(file_id):
    """
    Endpoint to remove a specific document from the collection.
    Note: This requires rebuilding the vector store without the removed document.
    """
    global document_metadata, vector_store, qa_chain, all_documents
    
    if file_id not in document_metadata:
        return jsonify({"error": "Document not found"}), 404
    
    try:
        # Get filename before removing
        filename = document_metadata[file_id]['filename']
        
        # Remove from metadata
        del document_metadata[file_id]
        
        # Remove uploaded file
        for file in os.listdir(UPLOAD_FOLDER):
            if file.startswith(file_id):
                os.remove(os.path.join(UPLOAD_FOLDER, file))
        
        # Filter out chunks belonging to this document
        remaining_docs = [doc for doc in all_documents if doc.metadata.get('file_id') != file_id]
        
        # Rebuild vector store if there are remaining documents
        if remaining_docs:
            all_documents = remaining_docs
            create_vector_store(remaining_docs, collection_name="docuchat_multi", add_to_existing=False)
            initialize_qa_chain()
        else:
            # No documents left, clear everything
            vector_store = None
            qa_chain = None
            all_documents = []
        
        # Update session state
        save_session_state()
        
        logger.info(f"Document removed: {filename}")
        
        return jsonify({
            "message": f"Document '{filename}' removed successfully",
            "remaining_documents": len(document_metadata)
        }), 200
        
    except Exception as e:
        logger.error(f"Error removing document: {str(e)}")
        return jsonify({"error": f"Failed to remove document: {str(e)}"}), 500

@app.route('/chat/sessions', methods=['GET'])
def get_chat_sessions():
    """Endpoint to retrieve all chat sessions (sidebar list)."""
    try:
        sessions_data = load_chat_sessions()
        # Return sessions as a list, sorted by updated_at (newest first)
        sessions_list = list(sessions_data['sessions'].values())
        sessions_list.sort(key=lambda x: x['updated_at'], reverse=True)
        
        # Return minimal info for sidebar
        sidebar_sessions = [{
            "id": s['id'],
            "name": s['name'],
            "created_at": s['created_at'],
            "updated_at": s['updated_at'],
            "message_count": len(s['messages'])
        } for s in sessions_list]
        
        return jsonify({
            "sessions": sidebar_sessions,
            "active_session": sessions_data.get('active_session'),
            "count": len(sidebar_sessions)
        }), 200
    except Exception as e:
        logger.error(f"Error retrieving sessions: {str(e)}")
        return jsonify({"error": f"Failed to retrieve sessions: {str(e)}"}), 500

@app.route('/chat/session/<session_id>', methods=['GET'])
def get_session(session_id):
    """Endpoint to retrieve a specific chat session with all messages and load its vector store."""
    try:
        sessions_data = load_chat_sessions()
        if session_id not in sessions_data['sessions']:
            return jsonify({"error": "Session not found"}), 404
        
        # Set as active session
        sessions_data['active_session'] = session_id
        save_chat_sessions(sessions_data)
        
        # Load this session's vector store (if it has documents)
        session = sessions_data['sessions'][session_id]
        if session.get('documents') and len(session['documents']) > 0:
            logger.info(f"Loading vector store for session {session_id}")
            load_session_vector_store(session_id)
        else:
            logger.info(f"Session {session_id} has no documents")
        
        return jsonify({
            "session": session
        }), 200
    except Exception as e:
        logger.error(f"Error retrieving session: {str(e)}")
        return jsonify({"error": f"Failed to retrieve session: {str(e)}"}), 500

@app.route('/chat/session/new', methods=['POST'])
def new_session():
    """Endpoint to create a new chat session."""
    try:
        session_id = create_new_session()
        sessions_data = load_chat_sessions()
        return jsonify({
            "session": sessions_data['sessions'][session_id]
        }), 201
    except Exception as e:
        logger.error(f"Error creating session: {str(e)}")
        return jsonify({"error": f"Failed to create session: {str(e)}"}), 500

@app.route('/chat/session/<session_id>/rename', methods=['PUT'])
def rename_session(session_id):
    """Endpoint to rename a chat session."""
    try:
        data = request.get_json()
        if not data or 'name' not in data:
            return jsonify({"error": "New name required"}), 400
        
        new_name = data['name'].strip()
        if not new_name or len(new_name) > 100:
            return jsonify({"error": "Invalid name length"}), 400
        
        sessions_data = load_chat_sessions()
        if session_id not in sessions_data['sessions']:
            return jsonify({"error": "Session not found"}), 404
        
        sessions_data['sessions'][session_id]['name'] = new_name
        sessions_data['sessions'][session_id]['updated_at'] = datetime.now().isoformat()
        save_chat_sessions(sessions_data)
        
        logger.info(f"Renamed session {session_id} to: {new_name}")
        return jsonify({
            "message": "Session renamed successfully",
            "session": sessions_data['sessions'][session_id]
        }), 200
    except Exception as e:
        logger.error(f"Error renaming session: {str(e)}")
        return jsonify({"error": f"Failed to rename session: {str(e)}"}), 500

@app.route('/chat/session/<session_id>', methods=['DELETE'])
def delete_session(session_id):
    """Endpoint to delete a chat session."""
    try:
        sessions_data = load_chat_sessions()
        if session_id not in sessions_data['sessions']:
            return jsonify({"error": "Session not found"}), 404
        
        session_name = sessions_data['sessions'][session_id]['name']
        del sessions_data['sessions'][session_id]
        
        # Clear active session if it was the deleted one
        if sessions_data.get('active_session') == session_id:
            sessions_data['active_session'] = None
        
        save_chat_sessions(sessions_data)
        logger.info(f"Deleted session: {session_name}")
        
        return jsonify({
            "message": f"Session '{session_name}' deleted successfully"
        }), 200
    except Exception as e:
        logger.error(f"Error deleting session: {str(e)}")
        return jsonify({"error": f"Failed to delete session: {str(e)}"}), 500

@app.route('/clear', methods=['POST'])
def clear_database():
    """
    Endpoint to clear the vector database and reset the system.
    Useful for starting fresh with new documents.
    Can optionally preserve chat history.
    """
    global vector_store, qa_chain, document_metadata, all_documents
    
    try:
        data = request.get_json() or {}
        clear_history = data.get('clear_history', True)
        
        # Reset global variables
        vector_store = None
        qa_chain = None
        document_metadata = {}
        all_documents = []
        
        # Clear uploaded files
        for filename in os.listdir(UPLOAD_FOLDER):
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
        
        # Clear session state
        save_session_state()
        
        # Optionally clear chat sessions
        if clear_history:
            save_chat_sessions({"sessions": {}, "active_session": None})
            logger.info("Database and history cleared successfully")
            message = "Database and chat history cleared successfully"
        else:
            logger.info("Database cleared (history preserved)")
            message = "Database cleared successfully. Chat history preserved."
        
        return jsonify({
            "message": message
        }), 200
        
    except Exception as e:
        logger.error(f"Error clearing database: {str(e)}")
        return jsonify({"error": f"Failed to clear database: {str(e)}"}), 500

@app.route('/restore', methods=['POST'])
def restore_previous_session():
    """
    Endpoint to manually restore previous session.
    Reloads documents and vector store from saved state.
    """
    try:
        success = restore_session()
        if success:
            return jsonify({
                "message": "Session restored successfully",
                "documents": document_metadata
            }), 200
        else:
            return jsonify({
                "message": "No previous session found or session is empty"
            }), 404
    except Exception as e:
        logger.error(f"Error restoring session: {str(e)}")
        return jsonify({"error": f"Failed to restore session: {str(e)}"}), 500

if __name__ == '__main__':
    # Starting fresh - no session restoration (documents reset on restart)
    logger.info("Starting DocuChat AI Backend...")
    logger.info("→ Starting with fresh session (documents will be cleared)")
    
    # Run Flask app in development mode
    app.run(debug=True, host='0.0.0.0', port=5000)
